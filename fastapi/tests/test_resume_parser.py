import os
import sys
import tempfile
import threading
from http.server import ThreadingHTTPServer, SimpleHTTPRequestHandler
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.core.config import Settings
from app.schemas.resume import ParsedResumeResponse
from app.services.resume_parser import parse_resume_from_url


class QuietHandler(SimpleHTTPRequestHandler):
    def log_message(self, format, *args):
        return


@pytest.fixture
def local_docx_server():
    with tempfile.TemporaryDirectory() as temp_dir:
        doc_path = Path(temp_dir) / "resume.docx"
        document = Document()
        document.add_paragraph("Jane Doe")
        document.add_paragraph("Senior Software Engineer")
        document.save(doc_path)

        handler = lambda *args, **kwargs: QuietHandler(*args, directory=temp_dir, **kwargs)
        server = ThreadingHTTPServer(("127.0.0.1", 0), handler)
        thread = threading.Thread(target=server.serve_forever, daemon=True)
        thread.start()

        try:
            yield f"http://127.0.0.1:{server.server_address[1]}/resume.docx"
        finally:
            server.shutdown()
            server.server_close()


def test_parse_resume_from_url_uses_public_url_when_internal_host_is_unreachable(local_docx_server):
    settings = Settings(backend_internal_base_url="http://127.0.0.1:1", gemini_api_key="")

    response = parse_resume_from_url(local_docx_server, settings)

    assert response is not None
    assert isinstance(response, ParsedResumeResponse)
